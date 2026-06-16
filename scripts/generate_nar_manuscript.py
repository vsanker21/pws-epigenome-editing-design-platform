#!/usr/bin/env python3
"""Generate extensive NAR Methods manuscript with narrative Results and cross-referenced supplementary materials."""

from __future__ import annotations

import sys
from datetime import datetime, timezone
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

ROOT = Path(r"g:\Prader Willi Syndrome")
sys.path.insert(0, str(ROOT / "scripts"))

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import manuscript_data as D
import manuscript_legends as L
import manuscript_methods_extensive as M

FIG_DIR = ROOT / "manuscript" / "figures"
OUT_DOC = ROOT / "manuscript" / "PWS_NAR_Methods_Manuscript.docx"


# ---------------------------------------------------------------------------
# Figure generation
# ---------------------------------------------------------------------------

def generate_all_figures() -> dict[str, Path]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    p: dict[str, Path] = {}

    # --- Main Figure 1: Pipeline ---
    fig, ax = plt.subplots(figsize=(13, 5.5))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 5.5)
    ax.axis("off")
    boxes = [
        (0.2, 2.8, "GEO + CRISPRepi\n23,506 records\n464 sig. sites"),
        (2.1, 2.8, "Digital twin\n1,675 nodes\n2,315 edges"),
        (4.0, 2.8, "Forward model\n57 Tet1 / 17 VP64\nrules score"),
        (5.9, 2.8, "Hybrid opt.\n121 designs\n96 dual-window"),
        (7.8, 2.8, "Validation\n4 datasets\noff-target + DMR"),
        (9.7, 2.8, "Catalog\n25 ranked\ncross-validated"),
        (11.6, 2.8, "Protocol\nwet-lab arms\nA–D"),
    ]
    colors = ["#4C72B0", "#55A868", "#C44E52", "#8172B2", "#CCB974", "#64B5CD", "#E17C05"]
    for i, (x, y, txt) in enumerate(boxes):
        rect = mpatches.FancyBboxPatch(
            (x, y), 1.55, 1.5, boxstyle="round,pad=0.05",
            facecolor=colors[i], edgecolor="black", alpha=0.88, linewidth=1.2,
        )
        ax.add_patch(rect)
        ax.text(x + 0.77, y + 0.75, txt, ha="center", va="center", fontsize=7.5, color="white", fontweight="bold")
        if i < len(boxes) - 1:
            ax.annotate("", xy=(x + 1.58, y + 0.75), xytext=(x + 2.0, y + 0.75),
                        arrowprops=dict(arrowstyle="->", lw=1.8, color="#333"))
    ax.set_title("Figure 1. End-to-end workflow from public data to ranked therapeutic designs", fontsize=11, fontweight="bold")
    p["fig1"] = _save(fig, "fig1_pipeline.png")

    # --- Main Figure 2: Editor failure modes ---
    fig, axes = plt.subplots(1, 2, figsize=(10, 4.5))
    genes = ["SNRPN", "SNHG14"]
    for ax, vals, title, color in [
        (axes[0], [52.3, 99.6], "dCas9-Tet1 (PWS-ICR demethylation)", "#4C72B0"),
        (axes[1], [246.3, 0.0], "dCas9-VP64 (SNRPN promoter activation)", "#C44E52"),
    ]:
        bars = ax.bar(genes, vals, color=color, width=0.55, edgecolor="black")
        ax.axhspan(70, 130, alpha=0.15, color="green", label="Therapeutic window")
        ax.set_ylabel("% WT expression (GSE243185)")
        ax.set_title(title, fontsize=10)
        ax.set_ylim(0, 280)
        for bar, v in zip(bars, vals):
            ax.text(bar.get_x() + bar.get_width() / 2, v + 6, f"{v:.1f}%", ha="center", fontsize=9)
        ax.legend(fontsize=7)
    fig.suptitle("Figure 2. Non-overlapping editor outcomes motivate hybrid design (Results §2)", fontsize=11, fontweight="bold")
    fig.tight_layout()
    p["fig2"] = _save(fig, "fig2_editor_outcomes.png")

    # --- Main Figure 3: Digital twin subregion map ---
    fig, ax = plt.subplots(figsize=(10, 3.5))
    subs = D.integration["subregion_hg38_distribution"]
    labels = list(subs.keys())
    vals = list(subs.values())
    colors_bar = ["#C44E52" if "icr" in l else "#4C72B0" for l in labels]
    ax.barh(labels, vals, color=colors_bar, edgecolor="black")
    ax.set_xlabel("Significant gRNA sites (n=464)")
    ax.set_title("Figure 3. Digital twin: 80.4% of significant hits target PWS-ICR (Results §1; Supp. Fig. S1)", fontsize=10, fontweight="bold")
    for i, v in enumerate(vals):
        ax.text(v + 3, i, str(v), va="center", fontsize=9)
    fig.tight_layout()
    p["fig3"] = _save(fig, "fig3_subregion_distribution.png")

    # --- Main Figure 4: Tet1 ICR proximity vs score ---
    tet1 = D.tet1_ranked[D.tet1_ranked["editor_system"] == "dCas9-Tet1"].copy()
    fig, ax = plt.subplots(figsize=(7, 5))
    ax.scatter(tet1["methylation_dist_bp"], tet1["rules_reactivation_score"],
               c=-np.log10(tet1["padj"].clip(1e-300)), cmap="viridis", s=40, edgecolors="k", linewidths=0.3)
    for gid, color in [("PWS_G.4670", "red"), ("PWS_G.4363", "orange")]:
        row = tet1[tet1["grna_id"] == gid]
        if len(row):
            ax.scatter(row["methylation_dist_bp"], row["rules_reactivation_score"], s=120, c=color,
                       edgecolors="black", zorder=5, label=gid)
    ax.set_xlabel("Distance to nearest assayed CpG (bp)")
    ax.set_ylabel("rules_reactivation_score")
    ax.set_title("Figure 4. ICR-proximal Tet1 guides score highest (ρ=−0.86; Results §2; Supp. Table S5)", fontsize=10, fontweight="bold")
    ax.legend(fontsize=8)
    fig.tight_layout()
    p["fig4"] = _save(fig, "fig4_tet1_distance_score.png")

    # --- Main Figure 5: Hybrid dose-response ---
    w = np.linspace(0, 0.3, 120)
    snrpn = (1 - w) * 52.3 + w * 246.3
    snhg14 = (1 - w) * 99.6 + w * 0.0
    fig, ax = plt.subplots(figsize=(7, 5))
    ax.plot(w, snrpn, "b-", lw=2.5, label="SNRPN")
    ax.plot(w, snhg14, "r-", lw=2.5, label="SNHG14")
    ax.axhspan(70, 130, alpha=0.12, color="green")
    ax.axvline(0.062, color="gray", ls="--", label="w_VP64=0.062 (rank 1)")
    ax.scatter([0.062], [70.0], s=140, c="gold", edgecolors="black", zorder=5)
    ax.scatter([0], [52.3], s=80, c="#4C72B0", edgecolors="black", zorder=5, label="Tet1 mono.")
    ax.set_xlabel("Relative VP64 dose (w_VP64)")
    ax.set_ylabel("Predicted % WT")
    ax.set_title("Figure 5. Hybrid optimization lands at SNRPN window boundary (Results §3; Supp. Fig. S4)", fontsize=10, fontweight="bold")
    ax.legend(fontsize=8)
    ax.set_xlim(0, 0.25)
    fig.tight_layout()
    p["fig5"] = _save(fig, "fig5_hybrid_dose.png")

    # --- Main Figure 6: Held-out validation ---
    guides = ["PWS_G.4670", "PWS_G.4363", "PWS_G.9414", "PWS_G.8738"]
    pct = [98.2, 61.4, 94.1, 88.2]
    fig, ax = plt.subplots(figsize=(8, 4.5))
    colors_bar = ["#55A868" if x >= 90 else "#C44E52" for x in pct]
    ax.barh(guides, pct, color=colors_bar, edgecolor="black")
    ax.axvline(90, color="black", ls="--", label="90th percentile")
    ax.set_xlabel("Within-editor percentile")
    ax.set_title("Figure 6. Retrospective guide recovery: 3/4 pass (Results §4; Supp. Table S1)", fontsize=10, fontweight="bold")
    for i, v in enumerate(pct):
        ax.text(v + 1, i, f"{v:.1f}%", va="center", fontsize=9)
    ax.legend(fontsize=8)
    fig.tight_layout()
    p["fig6"] = _save(fig, "fig6_held_out.png")

    # --- Main Figure 7: Multi-dataset validation story ---
    fig, axes = plt.subplots(1, 3, figsize=(12, 4))
    # A organoid
    ax = axes[0]
    ax.bar(["SNRPN", "SNHG14"], [5297, 158.1], color="#8172B2", edgecolor="black")
    ax.set_yscale("log")
    ax.set_ylabel("Fold-change (organoid)")
    ax.set_title("A. Organoid (GSE262700)")
    # B neuron vs iPSC Tet1
    ax = axes[1]
    x = np.arange(2)
    ax.bar(x - 0.2, [52.3, 99.6], 0.4, label="iPSC bulk", color="#4C72B0")
    ax.bar(x + 0.2, [47.2, 56.8], 0.4, label="Neuron (GSE285305)", color="#64B5CD")
    ax.set_xticks(x)
    ax.set_xticklabels(["SNRPN", "SNHG14"])
    ax.set_ylabel("% WT")
    ax.set_title("B. Cell-type concordance")
    ax.legend(fontsize=7)
    # C ICR bigWig
    ax = axes[2]
    ct = ["ESC", "Hyp. prog.", "Hyp. neuron"]
    sig = [0.057, 0.332, 0.603]
    ax.bar(ct, sig, color=["#CCB974", "#64B5CD", "#C44E52"], edgecolor="black")
    ax.set_ylabel("ICR bigWig signal")
    ax.set_title(f"C. GSE152098 ({D.gse152098['neuron_vs_esc_icr_fold']:.1f}× neuron/ESC)")
    fig.suptitle("Figure 7. Convergent external validation across three independent datasets (Results §4–5)", fontsize=10, fontweight="bold")
    fig.tight_layout()
    p["fig7"] = _save(fig, "fig7_multidataset_validation.png")

    # --- Main Figure 8: Safety ---
    fig, axes = plt.subplots(1, 2, figsize=(10, 4))
    ax = axes[0]
    ax.bar(["PWS_G.4670\n(Tet1)", "PWS_G.9414\n(VP64)"], [1.438, 2.625],
           color=["#4C72B0", "#C44E52"], edgecolor="black")
    ax.set_ylabel("Cas-OFFinder-compatible score")
    ax.set_title("A. Genome-wide off-target")
    ax = axes[1]
    imprinted = D.gse28525[D.gse28525["imprinted_dmr"] == True]
    icr_probes = imprinted[imprinted["maternally_methylated"] == True]
    ax.bar(["Imprinted\nDMRs", "ICR maternal\nprobes", "GSE298378\nprobes"],
           [D.dmr_map["gse28525"]["imprinted_dmrs_in_window"],
            D.dmr_map["gse28525"]["icr_imprinted_dmrs"],
            D.dmr_map["gse298378"]["probes_in_window"]],
           color="#55A868", edgecolor="black")
    ax.set_ylabel("Count in PWS window")
    ax.set_title("B. Collateral methylation map")
    fig.suptitle("Figure 8. Safety profiling integrates off-target scan and imprinting DMR context (Results §6; Supp. Figs S7–S8)", fontsize=10, fontweight="bold")
    fig.tight_layout()
    p["fig8"] = _save(fig, "fig8_safety.png")

    # --- Supplementary figures ---
    # S1: Digital twin node composition
    fig, ax = plt.subplots(figsize=(7, 4))
    nc = D.integration["node_counts"]
    ax.pie(nc.values(), labels=[f"{k}\n({v})" for k, v in nc.items()], autopct="%1.1f%%", startangle=140)
    ax.set_title("Supplementary Figure S1. Digital twin node composition (Methods §2.2; Results §1)")
    p["sfig1"] = _save(fig, "supp_fig_s1_digital_twin_nodes.png")

    # S2: Monte Carlo sensitivity
    fig, ax = plt.subplots(figsize=(7, 4))
    top5 = D.sensitivity["top5"]
    ax.bar([str(d["catalog_rank"]) for d in top5], [d["p_both_in_window"] for d in top5],
           color="#8172B2", edgecolor="black")
    ax.axhline(0.5, color="red", ls="--")
    ax.set_xlabel("Catalog rank")
    ax.set_ylabel("P(both genes in 70–130% window)")
    ax.set_title("Supplementary Figure S2. Uncertainty analysis (Results §3; Methods §2.6)")
    p["sfig2"] = _save(fig, "supp_fig_s2_sensitivity.png")

    # S3: Pareto frontier
    pareto_opt = D.pareto_df[D.pareto_df["pareto_rank"].notna()].head(10)
    fig, ax = plt.subplots(figsize=(6, 5))
    ax.scatter(D.pareto_df["obj_SNRPN"], D.pareto_df["obj_SNHG14"], alpha=0.3, s=20, c="gray", label="All designs")
    if len(pareto_opt):
        ax.scatter(pareto_opt["obj_SNRPN"], pareto_opt["obj_SNHG14"], s=80, c="#C44E52", edgecolors="black", label="Pareto optimal")
    ax.set_xlabel("SNRPN window objective")
    ax.set_ylabel("SNHG14 window objective")
    ax.set_title(f"Supplementary Figure S3. Pareto frontier (n={D.pareto['n_pareto_optimal']}/121; Results §3; Methods §2.5)")
    ax.legend(fontsize=8)
    fig.tight_layout()
    p["sfig3"] = _save(fig, "supp_fig_s3_pareto.png")

    # S4: Bayesian GP samples
    fig, ax = plt.subplots(figsize=(6, 5))
    wv = D.gp_samples["w_vp64"]
    snrpn_gp = (1 - wv) * 52.3 + wv * 246.3
    sc = ax.scatter(wv, snrpn_gp, c=D.gp_samples["objective"], cmap="plasma", s=40, alpha=0.8)
    plt.colorbar(sc, ax=ax, label="Objective")
    ax.scatter([D.bayesian["gp_best_ei"]["w_vp64"]], [D.bayesian["gp_best_ei"]["predicted_SNRPN_pct"]],
               s=150, c="red", marker="*", edgecolors="black", label="GP best EI")
    ax.scatter([0.062], [70.0], s=100, c="gold", edgecolors="black", label="DE optimum (G4670)")
    ax.set_xlabel("w_VP64")
    ax.set_ylabel("Predicted SNRPN % WT")
    ax.set_title("Supplementary Figure S4. Bayesian GP dose surface (Methods §2.5; Supp. Table S4)")
    ax.legend(fontsize=7)
    fig.tight_layout()
    p["sfig4"] = _save(fig, "supp_fig_s4_bayesian_gp.png")

    # S5: ODE circuit scenarios
    fig, ax = plt.subplots(figsize=(8, 4.5))
    scenarios = D.circuit_sim["scenario"].tolist()
    hp = D.circuit_sim["hyperphagia_proxy"].tolist()
    colors_ode = ["#C44E52" if s == "UPD15_untreated" else "#55A868" if s == "WT_reference" else "#4C72B0" for s in scenarios]
    ax.barh(scenarios, hp, color=colors_ode, edgecolor="black")
    ax.set_xlabel("Hyperphagia proxy (normalized)")
    ax.set_title("Supplementary Figure S5. ODE circuit exploration (Methods §2.7; not used for ranking)")
    fig.tight_layout()
    p["sfig5"] = _save(fig, "supp_fig_s5_ode_circuit.png")

    # S6: Organoid fold changes all genes
    fig, ax = plt.subplots(figsize=(8, 4))
    ofc = D.organoid_fc.sort_values("fold_change_edited_vs_control", ascending=True)
    ax.barh(ofc["gene_symbol"], ofc["fold_change_edited_vs_control"].clip(0.01, 10000), color="#8172B2", edgecolor="black")
    ax.set_xscale("log")
    ax.set_xlabel("Fold-change (edited vs control organoid)")
    ax.set_title("Supplementary Figure S6. Organoid reactivation landscape (Results §4; GSE262700)")
    fig.tight_layout()
    p["sfig6"] = _save(fig, "supp_fig_s6_organoid_fc.png")

    # S7: Cas-OFFinder mismatch counts
    fig, ax = plt.subplots(figsize=(7, 4))
    for gid, color, label in [("PWS_G.4670", "#4C72B0", "Tet1"), ("PWS_G.9414", "#C44E52", "VP64")]:
        counts = D.cas_off["per_guide"][gid]["counts"]
        mm = [counts[f"mm{i}"] for i in range(5)]
        ax.plot(range(5), mm, "o-", color=color, label=f"{label} ({gid})", lw=2)
    ax.set_xticks(range(5))
    ax.set_xticklabels(["mm0", "mm1", "mm2", "mm3", "mm4"])
    ax.set_ylabel("Genome hit count")
    ax.set_title("Supplementary Figure S7. Cas-OFFinder-compatible mismatch spectrum (Results §6; Methods §2.4)")
    ax.legend(fontsize=8)
    fig.tight_layout()
    p["sfig7"] = _save(fig, "supp_fig_s7_cas_offinder_mm.png")

    # S8: GSE28525 ICR probe methylation beta
    icr_probes = D.gse28525[(D.gse28525["imprinted_dmr"] == True) & (D.gse28525["maternally_methylated"] == True)]
    if len(icr_probes):
        fig, ax = plt.subplots(figsize=(8, 4))
        tissues = ["beta_pUPD_mean", "beta_mUPD", "beta_Brain"]
        means = [icr_probes[t].mean() for t in tissues]
        ax.bar(["Paternal UPD", "Maternal UPD", "Brain ref."], means, color=["#4C72B0", "#C44E52", "#55A868"], edgecolor="black")
        ax.set_ylabel("Mean beta (450K)")
        ax.set_title("Supplementary Figure S8. ICR imprinted DMR methylation reference (Results §6; Methods §2.3)")
        fig.tight_layout()
        p["sfig8"] = _save(fig, "supp_fig_s8_icr_methylation.png")

    # S9: Neuron vs iPSC delta
    fig, ax = plt.subplots(figsize=(6, 4))
    comps = D.neuron_conc["comparisons"]
    labels = [c["editor_gene"].split(":")[1] for c in comps]
    deltas = [c["delta_pct"] for c in comps]
    colors_n = ["#55A868" if c["direction_concordant"] else "#C44E52" for c in comps]
    ax.barh(labels, deltas, color=colors_n, edgecolor="black")
    ax.axvline(0, color="black", lw=0.8)
    ax.set_xlabel("Δ % WT (neuron − iPSC)")
    ax.set_title("Supplementary Figure S9. Neuron-specific RNA concordance (Results §4; GSE285305)")
    fig.tight_layout()
    p["sfig9"] = _save(fig, "supp_fig_s9_neuron_concordance.png")

    # S10: Catalog uncertainty vs off-target
    fig, ax = plt.subplots(figsize=(6, 5))
    c = D.catalog[D.catalog["strategy"] == "Hybrid_Tet1_VP64"]
    ax.scatter(c["uncertainty"], c["genome_offtarget_risk"], c=c["catalog_rank"], cmap="viridis", s=60, edgecolors="k")
    ax.set_xlabel("Design uncertainty")
    ax.set_ylabel("Genome off-target risk score")
    ax.set_title("Supplementary Figure S10. Catalog safety–confidence landscape (Results §3; Supp. Table S6)")
    fig.tight_layout()
    p["sfig10"] = _save(fig, "supp_fig_s10_catalog_landscape.png")

    return p


def _save(fig, name: str) -> Path:
    path = FIG_DIR / name
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def set_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15


def fig_block(doc: Document, path: Path, legend: str, width: float = 6.2) -> None:
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, block in enumerate(legend.split("\n\n")):
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = cap.add_run(block.strip())
        r.italic = True
        r.font.size = Pt(10)
        if i == 0:
            r.bold = True
        cap.paragraph_format.space_after = Pt(8)
    doc.add_paragraph()


def table_legend_block(doc: Document, legend: str) -> None:
    """Detailed table legend and inference below table."""
    for i, block in enumerate(legend.split("\n\n")):
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = cap.add_run(block.strip())
        r.italic = True
        r.font.size = Pt(10)
        if i == 0:
            r.bold = True
        cap.paragraph_format.space_after = Pt(8)


def table_block(doc: Document, caption: str, headers: list[str], rows: list[list]) -> None:
    cap = doc.add_paragraph()
    r = cap.add_run(caption)
    r.bold = True
    r.font.size = Pt(10)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Manuscript body
# ---------------------------------------------------------------------------

def build_document(fig: dict[str, Path]) -> None:
    doc = Document()
    set_style(doc)
    T = D.TOP
    main_legends = L.main_figure_legends()
    supp_legends = L.supplementary_figure_legends()
    table_notes = L.main_table_notes()
    supp_table_notes = L.supplementary_table_notes()

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(
        "An allele-aware computational design platform for CRISPR epigenome-editing "
        "therapy in maternal UPD15 Prader–Willi syndrome"
    )
    tr.bold = True
    tr.font.size = Pt(16)
    for line in ["[Authors]", "[Affiliations]", "Correspondence: [email]",
                 f"Manuscript generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}",
                 "Article type: Methods (Nucleic Acids Research)"]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ABSTRACT
    heading(doc, "Abstract")
    para(doc, (
        "Maternal uniparental disomy of chromosome 15 (UPD15) causes Prader–Willi syndrome (PWS) through "
        "biallelic silencing of paternally expressed genes at 15q11–q13. Genome-scale CRISPR epigenome-editing "
        "screens now demonstrate reactivation, yet no framework exists to prioritize allele-aware, multi-editor "
        "therapeutic configurations from complementary mechanism classes. We built a computational design platform "
        "integrating a GRCh38 PWS-locus digital twin (1,675 nodes; 2,315 edges), editor-stratified forward models "
        "(57 dCas9-Tet1 and 17 dCas9-VP64 significant guides), hybrid dose optimization, and multi-layer "
        "retrospective validation. From 23,506 curated editing records (464 significant sites), differential evolution "
        "evaluated 121 hybrid designs; 96 achieved dual-gene expression within a 70–130% wild-type therapeutic window. "
        "The top-ranked hybrid combines PWS_G.4670 (Tet1, chr15:24,954,921) with PWS_G.9414 (VP64, chr15:24,862,855) "
        "at w_Tet1=0.987 and w_VP64=0.062, predicting SNRPN at 70.0% and SNHG14 at 98.3% of WT. Cross-editor machine "
        "learning was explicitly rejected (CV R²<0). Held-out validation recovered 3/4 experimentally validated guides "
        "in the top decile within editor class (Tet1 Spearman ρ=0.401, p=0.002; VP64 ρ=0.630, p=0.007). Organoid "
        "scRNA-seq (GSE262700), neuron RNA (GSE285305), hypothalamic ATAC bigWigs (GSE152098), 450K methylation maps "
        "(GSE28525/GSE298378), and Cas-OFFinder-compatible genome scans convergently support biological plausibility "
        "and safety context. Monte Carlo analysis yields p(both genes in window)=0.498 for the top design. We provide "
        "a prospective experimental protocol for UPD15 iPSC-derived hypothalamic neurons. This platform is a "
        "design-prioritization layer atop published reactivation evidence, not a therapeutic efficacy demonstration."
    ))
    para(doc, "Keywords: ", bold=True)
    para(doc, "Prader–Willi syndrome; imprinting; CRISPR epigenome editing; computational therapeutic design; "
              "dCas9-Tet1; dCas9-VP64; 15q11-q13; digital twin")

    heading(doc, "Significance Statement", 2)
    para(doc, (
        "Therapeutic epigenome editing for imprinting disorders requires mechanism-aware design beyond standard "
        "CRISPR off-target tools. We introduce the first integrated platform that maps UPD15 PWS biology onto a "
        "unified genomic model, derives a testable hybrid Tet1+VP64 hypothesis from complementary editor failure modes, "
        "and ranks 25 uncertainty-aware designs validated against four public datasets—with an experimental protocol "
        "for wet-lab confirmation."
    ))

    doc.add_page_break()

    # INTRODUCTION (condensed but connected)
    heading(doc, "Introduction")
    para(doc, (
        "Prader–Willi syndrome (PWS; OMIM #176270) arises when paternally expressed genes at 15q11–q13 remain "
        "silent on both homologues. In ~25% of cases, maternal UPD15 imposes a fully maternal epigenotype at the "
        "PWS imprinting control region (ICR), silencing SNRPN and the SNHG14/SNORD116 transcript cluster that "
        "underlies neurobehavioral phenotypes including hyperphagia. Two recent studies established biological "
        "feasibility: Nemoto et al. (2025) achieved massive SNRPN/SNHG14 reactivation in PWS hypothalamic organoids "
        "via SunTag-TET1 ICR demethylation (GSE262700); Rohm et al. (2025) screened dCas9-Tet1 and dCas9-VP64 gRNAs "
        "in UPD15 iPSCs, revealing that Tet1 restores SNHG14 (~99.6% WT) but under-restores SNRPN (~52.3% WT), "
        "whereas VP64 hyperactivates SNRPN (~246.3% WT) without rescuing SNHG14 (GSE285306, GSE243185)."
    ))
    para(doc, (
        "These findings answer whether reactivation is possible; they do not answer which guide combinations and "
        "editor dose ratios achieve dual-gene paternal-equivalent expression with acceptable collateral imprinting risk. "
        "Standard CRISPR design tools optimize nuclease efficiency; CRISPRepi catalogs epigenome-editing records but "
        "provides no locus-specific therapeutic prioritization. Here we present an allele-aware computational platform "
        "that closes this gap by weaving five analytical layers into a single narrative: (i) a multi-omic digital twin, "
        "(ii) editor-stratified forward modeling with rejected cross-editor ML, (iii) mechanism-derived hybrid optimization, "
        "(iv) convergent retrospective validation across independent datasets, and (v) an uncertainty-aware ranked catalog "
        "with prospective experimental protocol (Figure 1; Supplementary Methods S1–S10)."
    ))

    # METHODS — extensive NAR-style section
    doc.add_page_break()
    M.write_materials_and_methods(doc)

    # RESULTS — narrative, extensive
    doc.add_page_break()
    heading(doc, "Results")

    heading(doc, "A unified digital twin reveals ICR-centric editing architecture at 15q11–q13", 2)
    para(doc, (
        "We first constructed an allele-aware digital twin of chr15:24,800,000–32,700,000 (GRCh38) by integrating "
        f"{D.locus_summary['gene_count']} annotated genes, {D.locus_summary['cpg_island_count']} CpG islands, "
        f"{D.integration['node_counts']['atac_peak']} hypothalamic ATAC peaks, {D.integration['node_counts']['methylation_cpg']} "
        f"assayed methylation CpGs, and {D.integration['merged_grna_sites']} base-pair-merged significant gRNA targets "
        f"into a graph of {D.integration['node_counts']['grna_target'] + sum(v for k,v in D.integration['node_counts'].items() if k != 'grna_target')} "
        f"nodes and {D.integration['edge_count']} edges (Figure 1; Figure 3; Supplementary Figure S1; Supplementary Methods S2)."
    ))
    para(doc, (
        f"The subregion distribution was sharply ICR-centric: {D.integration['subregion_hg38_distribution']['pws_icr']} of 464 "
        f"significant sites ({100*D.integration['subregion_hg38_distribution']['pws_icr']/464:.1f}%) mapped to the PWS-ICR, "
        f"compared with {D.integration['subregion_hg38_distribution']['snrpn_snh14']} at the SNRPN–SNHG14 transcriptional unit "
        f"and only {D.integration['subregion_hg38_distribution']['outside_pws_critical']} outside the critical region. "
        "This spatial architecture is consistent with Rohm et al.'s finding that ICR demethylation is the dominant route "
        "to SNHG14 rescue, and it constrains therapeutic design to a compact ~7 kb ICR window (hg38 chr15:24,948,846–24,955,292) "
        "plus a discrete SNRPN promoter VP64 target class ~92 kb telomeric (Supplementary Table S7)."
    ))
    para(doc, (
        f"Hypothalamic relevance was quantified by reprocessing 14 GSE152098 bigWig files. Mean signal at the PWS-ICR was "
        f"{D.gse152098['icr_celltype_means'][2]['icr_mean_signal']:.3f} in hypothalamic neurons versus "
        f"{D.gse152098['icr_celltype_means'][0]['icr_mean_signal']:.3f} in ESCs "
        f"({D.gse152098['neuron_vs_esc_icr_fold']:.1f}-fold; hypothalamic progenitors intermediate at 0.332; "
        "Figure 7C; Supplementary Methods S3). ENCODE regulatory tracks (8/8 successful; 5,000 features per track via UCSC API) "
        "provided brain/embryo DNase, H3K27ac, H3K4me3, CTCF, and TFBS context (Supplementary Table S3; Supplementary Methods S4)."
    ))
    fig_block(doc, fig["fig1"], main_legends["fig1"])
    fig_block(doc, fig["fig3"], main_legends["fig3"])

    heading(doc, "Editor-stratified forward modeling exposes complementary gaps that monotherapy cannot close", 2)
    para(doc, (
        "The biological story hinges on a simple but consequential asymmetry: Tet1 and VP64 are not interchangeable editors "
        "but mechanism-complementary tools acting on different cis-regulatory elements. At the editor-class level (GSE243185), "
        "dCas9-Tet1 yielded SNRPN at 52.3% and SNHG14 at 99.6% of WT iPSC expression; dCas9-VP64 yielded SNRPN at 246.3% "
        "and SNHG14 at 0% (Figure 2). No single editor places both genes within the predefined 70–130% therapeutic window."
    ))
    para(doc, (
        "We therefore rejected cross-editor machine learning (pooled Tet1+VP64 training set: 5-fold CV R²<0) and deployed "
        "editor-stratified rules_reactivation_score as the primary ranking function (Supplementary Methods S5). For Tet1, "
        f"distance to the ICR anchor correlated with score (Spearman ρ={D.held_out['tet1_distance_benchmark']['dist_vs_score_spearman']:.2f}, "
        f"n=57) and methylation change correlated positively (ρ={D.held_out['tet1_distance_benchmark']['methylation_vs_score_spearman']:.3f}, "
        f"n=48 guides with methylation data). The top Tet1 guide PWS_G.4670 (protospacer GCAGGCTGGCGCGCATGCTC, chr15:24,954,921) "
        f"achieved rules_reactivation_score=1.02 (padj=9.2×10⁻¹⁵) and edited ICR methylation of 37.2% (Δ=−54.9% vs non-targeting; "
        "Supplementary Table S5). A supplementary within-editor Ridge model (features: log distance, SNHG14 targeting flag) "
        f"achieved CV R²={D.forward['tet1_within_editor_model']['cv_r2_mean']:.2f} and is reported for transparency only."
    ))
    para(doc, (
        "Critically, screen ranking and experimental methylation validation capture partially orthogonal properties. "
        f"PWS_G.4363—bisulfite-validated in GSE285300 (padj=2.95×10⁻⁶)—ranked 22nd of 57 Tet1 guides (61.4th percentile; "
        "Figure 6). We retain G4363 as an experimental methylation comparator rather than dismissing it for sub-decile rank "
        "(Supplementary Table S1; Discussion). VP64 guides PWS_G.9414 and PWS_G.8738 ranked 1st and 2nd within their class "
        f"(padj=7.5×10⁻¹³³ and 7.77×10⁻⁹³)."
    ))
    fig_block(doc, fig["fig2"], main_legends["fig2"])
    fig_block(doc, fig["fig4"], main_legends["fig4"])

    heading(doc, "Hybrid optimization identifies a minimal VP64 supplementation regime at the SNRPN window boundary", 2)
    para(doc, (
        "Given the dual-gap structure, we asked whether a phenomenological two-editor dose model could identify a combined "
        "regime placing both SNRPN and SNHG14 within the therapeutic window. Differential evolution over (w_Tet1, w_VP64) "
        f"evaluated {D.optimization['summary']['n_designs_evaluated']} guide-pair configurations; "
        f"{D.optimization['summary']['n_both_genes_in_target_window']} achieved dual-window outcomes "
        f"({100*D.optimization['summary']['n_both_genes_in_target_window']/D.optimization['summary']['n_designs_evaluated']:.1f}%). "
        "The optimizer consistently converged to near-maximal Tet1 dose with minimal VP64 supplementation "
        "(w_Tet1≈0.987, w_VP64≈0.062), landing SNRPN precisely at the 70% lower boundary while preserving SNHG14 at ~98% "
        "(Figure 5; Supplementary Figure S4; Supplementary Methods S6)."
    ))
    para(doc, (
        f"The top catalog entry pairs PWS_G.4670 with PWS_G.9414 (Table 1). Pareto non-dominated sorting identified "
        f"{D.pareto['n_pareto_optimal']} designs on the four-objective frontier (SNRPN window, SNHG14 window, safety, confidence; "
        f"Supplementary Figure S3). All Pareto-optimal hybrids included PWS_G.4670 as the Tet1 component with predicted "
        "70.0%/98.3% WT. Bayesian Gaussian Process optimization on an alternative pair (G4363+G9414) suggested higher VP64 "
        f"weighting (w_VP64≈{D.bayesian['gp_best_ei']['w_vp64']:.2f}; Supplementary Figure S4; Supplementary Table S4), "
        "confirming guide-pair dependence of the phenomenological surrogate."
    ))
    para(doc, (
        f"Uncertainty quantification tempered deterministic predictions. Monte Carlo analysis (50,000 samples per design; "
        f"σ=25×uncertainty) estimated p(SNRPN in window)={D.sensitivity['top_design']['p_SNRPN_in_window']:.3f}, "
        f"p(SNHG14 in window)={D.sensitivity['top_design']['p_SNHG14_in_window']:.1f}, and "
        f"p(both in window)={D.sensitivity['top_design']['p_both_in_window']:.3f} for the top design "
        "(Supplementary Figure S2; Supplementary Methods S8). SNHG14 predictions are robust; SNRPN sits at the window "
        "threshold under perturbation, mandating experimental VP64 titration rather than trusting point estimates alone."
    ))
    fig_block(doc, fig["fig5"], main_legends["fig5"])
    table_block(doc, "Table 1. Top-ranked hybrid therapeutic design (catalog rank 1).",
        ["Parameter", "Value"],
        [
            ["Strategy", "Hybrid_Tet1_VP64"],
            ["Tet1 guide", f"{T['tet1_grna_id']} — {T['tet1_protospacer']}"],
            ["Tet1 hg38", f"chr15:{int(T['tet1_hg38_start']):,}"],
            ["VP64 guide", f"{T['vp64_grna_id']} — {T['vp64_protospacer']}"],
            ["VP64 hg38", f"chr15:{int(T['vp64_hg38_start']):,}"],
            ["Dosing w_Tet1 / w_VP64", f"{T['recommended_w_tet1']:.3f} / {T['recommended_w_vp64']:.3f}"],
            ["Predicted SNRPN / SNHG14 (% WT)", f"{T['predicted_SNRPN_pct_WT']:.1f} / {T['predicted_SNHG14_pct_WT']:.1f}"],
            ["Objective score", f"{T['objective_score']:.2f}"],
            ["Uncertainty", f"{T['uncertainty']:.3f}"],
            ["Genome off-target risk", f"{T['genome_offtarget_risk']:.3f}"],
            ["p(both in 70–130% window)", f"{D.sensitivity['top_design']['p_both_in_window']:.3f}"],
            ["Experimental evidence", "Rohm 2025 top Tet1 + top VP64 hits"],
        ])
    table_legend_block(doc, table_notes["table1"])

    heading(doc, "Convergent retrospective validation links screen rankings to organoid, neuronal, and methylation readouts", 2)
    para(doc, (
        "A design platform earns credibility only if its rankings and predictions align with independent experiments. "
        f"We assembled a four-layer validation suite (Figure 6–7; Supplementary Tables S1, S8–S9). Layer 1—held-out guide "
        f"recovery—tested four experimentally validated guides against rules_reactivation_score within editor class. "
        f"Three of four ({100*D.held_out['validated_guide_recovery']['n_pass_top_decile']/D.held_out['validated_guide_recovery']['n_validated']:.0f}%) "
        f"exceeded the 90th percentile (Figure 6). Layer 2—organoid concordance—compared Nemoto 2025 scRNA-seq (GSE262700) "
        f"with Rohm bulk RNA. Both SNRPN and SNHG14 showed concordant upregulation; organoid fold-changes were "
        f"{D.organoid['direction_concordance']['genes'][0]['organoid_fold_change']:.0f}× (SNRPN) and "
        f"{D.organoid['direction_concordance']['genes'][1]['organoid_fold_change']:.0f}× (SNHG14) from near-zero baseline "
        "(Supplementary Figure S6). Magnitudes are not directly comparable across platforms, but direction concordance "
        "confirms Tet1-class editing reactivates both genes in hypothalamic tissue—the clinically relevant context."
    ))
    para(doc, (
        f"Layer 3—neuron-specific RNA (GSE285305)—showed Tet1-direction concordance for SNRPN (iPSC 52.3% vs neuron 47.2% WT) "
        f"and SNHG14 (99.6% vs 56.8% WT), with neuron SNHG14 attenuated (Supplementary Figure S9). VP64 SNRPN showed "
        "discordant magnitude (iPSC 246.3% vs neuron 25.7%), underscoring cell-type-specific dose requirements. "
        f"Layer 4—bisulfite and 450K methylation—mapped {D.dmr_map['gse28525']['probes_in_window']} GSE28525 probes and "
        f"{D.dmr_map['gse298378']['probes_in_window']} GSE298378 probes to hg38 PWS coordinates, identifying "
        f"{D.dmr_map['gse28525']['imprinted_dmrs_in_window']} imprinted DMRs including "
        f"{D.dmr_map['gse28525']['icr_imprinted_dmrs']} ICR-associated maternally methylated probes for collateral monitoring "
        "(Supplementary Figure S8; Supplementary Methods S7)."
    ))
    fig_block(doc, fig["fig6"], main_legends["fig6"])
    fig_block(doc, fig["fig7"], main_legends["fig7"])

    heading(doc, "Integrated safety assessment supports low off-target and collateral risk for ICR-targeted hybrids", 2)
    para(doc, (
        f"Therapeutic design must balance efficacy against off-target editing and imprinting collateral. At the locus level, "
        f"both recommended protospacers showed zero exact duplicates and zero 1-mm seed matches among 464 screened sites "
        f"(Supplementary Table S10). Genome-wide Cas-OFFinder-compatible scans across hg38 assigned low risk labels: "
        f"PWS_G.4670 score={D.cas_off['per_guide']['PWS_G.4670']['cas_offinder_score']:.3f} "
        f"(mm0=1, mm3=1, mm4=5); PWS_G.9414 score={D.cas_off['per_guide']['PWS_G.9414']['cas_offinder_score']:.3f} "
        f"(mm0=1, mm2=1, mm3=2, mm4=18; Supplementary Figure S7). Top hybrid designs target the ICR >3.5 Mb from the GABAA "
        "cluster (chr15:28.5–31.0 Mb), yielding zero GABAA proximity flags across all 25 catalog entries."
    ))
    para(doc, (
        f"Collateral bulk RNA under Tet1 showed UBE3A at 51.8% WT and ATP10A at 65.2% WT (low change); NDN/MKRN3 remained "
        "silenced as expected for paternal-only genes on the maternal chromosome (Supplementary Table S2). Combined assessment: "
        "low-to-moderate overall collateral risk, with VP64 dose minimization as the primary mitigation for SNRPN overshoot."
    ))
    fig_block(doc, fig["fig8"], main_legends["fig8"])

    heading(doc, "An experimental protocol operationalizes the top hybrid hypothesis for wet-lab validation", 2)
    para(doc, (
        "The computational story culminates in a falsifiable experimental plan (Supplementary Table S12; "
        "experimental_validation_protocol.json). Primary model: PWS UPD15 iPSC-derived hypothalamic neurons/organoids. "
        "Arm A tests hybrid PWS_G.4670+PWS_G.9414 at w_VP64≈0.06; Arm B tests Tet1 monotherapy; Arms C–D provide "
        "unedited PWS and healthy WT references. Primary endpoints: PWS-ICR bisulfite methylation (~40% target from ~95% baseline), "
        "SNRPN/SNHG14 expression (70–130% WT), and collateral panel (UBE3A, ATP10A, GABRB3). Delivery via RNP electroporation "
        "or LNP-mRNA for VP64 titration. Measured outcomes will recalibrate the forward model in a closed design–build–test loop."
    ))

    table_block(doc, "Table 2. Top 10 catalog designs (n=25 total).",
        ["Rank", "Tet1", "VP64", "SNRPN%", "SNHG14%", "Uncertainty", "Off-target"],
        [[str(int(r.catalog_rank)), r.tet1_grna_id,
          r.vp64_grna_id if pd.notna(r.vp64_grna_id) else "—",
          f"{r.predicted_SNRPN_pct_WT:.1f}", f"{r.predicted_SNHG14_pct_WT:.1f}",
          f"{r.uncertainty:.3f}", f"{r.genome_offtarget_risk:.2f}"]
         for r in D.catalog.head(10).itertuples()])
    table_legend_block(doc, table_notes["table2"])

    # DISCUSSION
    doc.add_page_break()
    heading(doc, "Discussion")
    para(doc, (
        "This study introduces a computational design-prioritization platform for CRISPR epigenome-editing therapy in "
        "maternal UPD15 PWS. The narrative arc—from ICR-centric digital twin, through complementary editor failure modes, "
        "to hybrid optimization and convergent validation—transforms thousands of screen hits into a single testable "
        "hypothesis: PWS_G.4670+PWS_G.9414 at 98.7%/6.2% editor weighting. We explicitly do not claim therapeutic efficacy; "
        "Nemoto 2025 and Rohm 2025 established reactivation feasibility. Our contribution is the design layer that asks "
        "which configuration achieves dual-gene paternal-equivalent expression with quantified uncertainty."
    ))
    para(doc, (
        "Several limitations define the boundary of inference. Editor-level bulk RNA priors cannot resolve guide-specific "
        "expression (exemplified by G4363). The hybrid dose-response is phenomenological. Held-out validation uses n=4 guides. "
        "Cas-OFFinder-compatible scans are in-silico. Graph propagation and ODE circuit models (Supplementary Methods S9–S10) "
        "were exploratory and did not influence catalog ranking. Despite these constraints, the platform reduces experimental "
        "search space by orders of magnitude and provides auditable, reproducible prioritization."
    ))

    # Data availability, refs, supplementary
    heading(doc, "Data Availability")
    para(doc, (
        "All source GEO/CRISPRepi data are public. Processed outputs listed in Supplementary Data Inventory "
        f"(Supplementary Table S11). Analysis code, configuration files, and processed data supporting catalog "
        f"generation are available at {D.GITHUB_REPO} and archived on Zenodo ({D.ZENODO_URL})."
    ))

    heading(doc, "References")
    for i, ref in enumerate([
        "Rohm J, et al. (2025) Cell Genomics 5:100770.",
        "Nemoto H, et al. (2025) Nat Commun (GSE262700).",
        "Shi X, et al. (2025) NAR 53(D1):D901.",
        "Cousminer DL, et al. (2021) Nat Commun (GSE152098).",
        "Bae S, et al. (2014) Bioinformatics 30:1473-1475.",
    ], 1):
        para(doc, f"{i}. {ref}")

    # SUPPLEMENTARY
    doc.add_page_break()
    heading(doc, "Supplementary Data")
    para(doc, "Supplementary Data are available at NAR online.")

    heading(doc, "Supplementary Methods", 1)
    _write_supplementary_methods(doc)

    heading(doc, "Supplementary Figures", 1)
    _write_supplementary_figures(doc, fig)

    heading(doc, "Supplementary Tables", 1)
    _write_supplementary_tables(doc)

    heading(doc, "Cross-reference index", 1)
    _write_cross_reference_index(doc)

    doc.save(str(OUT_DOC))
    print(f"Saved: {OUT_DOC}")


def _write_supplementary_methods(doc: Document) -> None:
    """Extended supplementary methods with algorithms, parameters, and cross-refs."""
    sections = [
        ("S1. Pipeline orchestration, versioning, and reproducibility", [
            "Master runner: scripts/run_pipeline.py. Valid phase arguments: 3, 5, 6, 9, 10, 11, 12, all.",
            "Phase 12 orchestrator: scripts/run_phase12.py (Cas-OFFinder, ENCODE, GSE152098, 450K DMR, catalog rebuild).",
            "Every output JSON includes generated_utc ISO-8601 timestamp. Catalog version: 1.0.",
            "Genome cache: data/genome/hg38/chr*.fa.gz (25 chromosomes).",
            "Re-execution: python scripts/run_pipeline.py --phase all regenerates all artifacts deterministically from public inputs.",
            "Cross-refs: Figure 1; Methods §2.1; Supplementary Table S11.",
        ]),
        ("S2. Digital twin graph construction (detailed)", [
            f"Node types and counts: {D.integration['node_counts']}.",
            "Edge construction rules: (i) gRNA→CpG island if Euclidean distance ≤500 bp; (ii) gRNA→gene if 20-mer overlap with gene body/promoter; (iii) gene↔gene if same subregion and ≤50 kb apart; (iv) gRNA→ATAC peak if overlap or ≤500 bp.",
            f"Total edges: {D.integration['edge_count']}. Significant gRNA nodes: padj<0.05 from GSE285306 after BH correction (Rohm et al.).",
            "Subregion classifier (hg38): pws_icr if start∈[24948846,24955292]; snrpn_snh14 if ∈[24750000,26300000] and not ICR; snord116_cluster if ∈[26300000,28500000]; gabaa_cluster if ∈[28500000,31000000].",
            f"Distribution: {D.integration['subregion_hg38_distribution']}.",
            "Output files: data/integrated/locus_merged_hg38.csv, data/integrated/integration_summary.json.",
            "Cross-refs: Figure 3; Supplementary Figure S1; Supplementary Table S7; Methods §2.4.",
        ]),
        ("S3. GSE152098 bigWig extraction and quantification", [
            "Source: GSE152098_RAW.tar via GEO. Extracted 14 bigWig files (Human ESC, hypothalamic progenitor, hypothalamic neuron replicates).",
            "Quantification: UCSC bigWigSummary (Linux binary via WSL; scripts/reprocess_gse152098_raw.py).",
            "Regions (hg38): pws_critical [24800000,32700000], pws_icr [24948846,24955292], snrpn_snh14 [24800000,26300000], snord116_cluster [26300000,28500000], gabaa_cluster [28500000,31000000].",
            f"ICR means: ESC={D.gse152098['icr_celltype_means'][0]['icr_mean_signal']:.4f}, Hyp.prog.={D.gse152098['icr_celltype_means'][1]['icr_mean_signal']:.4f}, Neuron={D.gse152098['icr_celltype_means'][2]['icr_mean_signal']:.4f}. Fold neuron/ESC={D.gse152098['neuron_vs_esc_icr_fold']:.2f}.",
            "Note: Windows paths with spaces required shlex.quote for WSL; dataPoints=1 returns single mean per region.",
            "Cross-refs: Figure 7C; Methods §2.4; Supplementary Table S8.",
        ]),
        ("S4. ENCODE/UCSC REST API query parameters", [
            "Base URL: https://api.genome.ucsc.edu/getData/track. Genome=hg38. Region=chr15:24800000-32700000.",
            "Tracks queried (8/8 success): encRegTfbsClustered, wgEncodeReg4DnaseAllBrain, wgEncodeReg4DnaseAllEmbryo, wgEncodeReg4MarkH3k27acAllBrain, wgEncodeReg4MarkH3k27acAllEmbryo, wgEncodeReg4MarkH3k4me3AllBrain, wgEncodeReg4MarkH3k4me3AllEmbryo, wgEncodeReg4MarkCtcfAllBrain.",
            "Limitation: maxItems=5000 per request (HTTP 206). Script: scripts/download_encode_reference.py.",
            "Cross-refs: Supplementary Table S3; Methods §2.5.",
        ]),
        ("S5. rules_reactivation_score: formula and parameters", [
            "Editor-stratified scoring. Cross-editor ML rejected: 5-fold CV R²<0 on pooled Tet1+VP64.",
            "Tet1 score components: (i) −log10(padj) weight 1.0; (ii) subregion bonus (pws_icr=1.0, snrpn_snh14=0.95); (iii) distance penalty log10(dist_to_ICR+1); (iv) methylation Δ weight if GSE285300 data available.",
            f"Validation correlations (n=57 Tet1): dist vs score ρ={D.held_out['tet1_distance_benchmark']['dist_vs_score_spearman']:.2f}; methylation vs score ρ={D.held_out['tet1_distance_benchmark']['methylation_vs_score_spearman']:.3f} (n=48).",
            "Rohm 2025 anchors (hard constraints on editor-class outcomes): Tet1 SNRPN=52.3%, SNHG14=99.6%; VP64 SNRPN=246.3%, SNHG14=0%.",
            f"CRISPRepi priors: Tet1 demethylation rate 62.5% (n=16); VP64 activation 60.4% (n=134).",
            f"Supplementary Ridge (Tet1): features=[log_dist, targets_snh14]; CV R²={D.forward['tet1_within_editor_model']['cv_r2_mean']:.2f}.",
            "Cross-refs: Figures 2, 4; Supplementary Table S5; Methods §2.6.",
        ]),
        ("S6. Hybrid optimization: differential evolution parameters", [
            "Optimizer: scipy.optimize.differential_evolution. Bounds: w_tet1∈[0,1], w_vp64∈[0,1], constraint w_tet1+w_vp64≤1.",
            "Phenomenological model: SNRPN(w)=(1−w)×52.3+w×246.3; SNHG14(w)=(1−w)×99.6.",
            f"Designs evaluated: {D.optimization['summary']['n_designs_evaluated']}. Dual-window: {D.optimization['summary']['n_both_genes_in_target_window']} ({100*D.optimization['summary']['n_both_genes_in_target_window']/D.optimization['summary']['n_designs_evaluated']:.1f}%).",
            "Collateral penalty: +1.0 if either guide within GABAA cluster [28500000,31000000] on hg38. All top designs: penalty=0.",
            "Primary optimum: w_tet1=0.987, w_vp64=0.062 → SNRPN=70.0%, SNHG14=98.3%.",
            "Cross-refs: Figure 5; Table 1; Supplementary Table S4; Methods §2.7.",
        ]),
        ("S7. 450K methylation probe mapping algorithm", [
            "Manifest: wgEncodeHaibMethyl450CpgIslandDetails.txt (15,259 chr15 probes after hg38 liftOver).",
            f"GSE28525: {D.dmr_map['gse28525']['probes_in_window']} probes in PWS window; {D.dmr_map['gse28525']['imprinted_dmrs_in_window']} imprinted DMRs; {D.dmr_map['gse28525']['icr_imprinted_dmrs']} ICR maternally methylated.",
            f"GSE298378: {D.dmr_map['gse298378']['probes_in_window']} probes; {D.dmr_map['gse298378']['n_samples']} samples; beta range {D.dmr_map['gse298378']['mean_beta_range']}.",
            "Imprinted DMR flag: |delta_mUPD|>0.2 AND maternally_methylated OR paternally_methylated per reference sample comparison.",
            "Script: scripts/map_methylation_dmrs.py. Cross-refs: Figure 8B; Supplementary Figure S8; Methods §2.4.",
        ]),
        ("S8. Monte Carlo uncertainty: sampling procedure", [
            f"n_samples={D.sensitivity['n_samples']:,} per design. Perturbation: SNRPN_obs~N(μ_SNRPN, σ), SNHG14_obs~N(μ_SNHG14, σ) with σ_pct=25×uncertainty.",
            f"Rank 1: uncertainty=0.147, σ=3.68%. p_SNRPN_in_window={D.sensitivity['top_design']['p_SNRPN_in_window']:.3f}, p_SNHG14={D.sensitivity['top_design']['p_SNHG14_in_window']:.1f}, p_both={D.sensitivity['top_design']['p_both_in_window']:.3f}.",
            "Limitations: Gaussian assumption; shared σ across genes; heuristic uncertainty not calibrated to biological replicates.",
            "Cross-refs: Supplementary Figure S2; Figure 5; Methods §2.11.",
        ]),
        ("S9. Graph propagation (exploratory)", [
            f"Algorithm: Personalized PageRank with damping α={D.graph['damping_alpha']}. Graph: {D.graph['n_nodes']} nodes, {D.graph['n_edges']} edges.",
            "Seed: gRNA target node(s). Diffusion to gene nodes (SNRPN, SNHG14, UBE3A, ATP10A).",
            "Result: SNRPN propagated ≈0 for all validated guides; SNHG14 ≈0.001. Not used for catalog ranking.",
            "Cross-refs: Discussion; data/models/graph_model/graph_propagation_report.json.",
        ]),
        ("S10. ODE circuit model (exploratory)", [
            "Three-node linear relaxation: dX/dt = (X_target − X)/τ with τ=1. Nodes: AgRP (orexigenic), POMC (anorexigenic), OXT.",
            "Input: rescue_fraction = mean(SNRPN%, SNHG14%)/100. WT baselines: AgRP=0.6, POMC=1.0, OXT=0.8.",
            "Scenarios: UPD15 untreated (5%/5% → hyperphagia_proxy=0.05); Tet1 mono (52.3%/99.6% → 0.760); hybrid rank1 (70.0%/98.3% → 0.842); WT (1.0).",
            "Not used for catalog ranking. Cross-refs: Supplementary Figure S5; data/models/circuit_ode/circuit_ode_simulations.csv.",
        ]),
        ("S11. Cas-OFFinder-compatible genome scan", [
            "Scripts: scripts/cas_offinder_scan.py, scripts/offtarget_core.py. PAM: NGG. Mismatches: 0–4 (mm0–mm4).",
            f"PWS_G.4670: score={D.cas_off['per_guide']['PWS_G.4670']['cas_offinder_score']:.3f}, counts={D.cas_off['per_guide']['PWS_G.4670']['counts']}.",
            f"PWS_G.9414: score={D.cas_off['per_guide']['PWS_G.9414']['cas_offinder_score']:.3f}, counts={D.cas_off['per_guide']['PWS_G.9414']['counts']}.",
            "Limitations: in-silico only; no bulges; no epigenetic modulation. CHANGE-seq/GUIDE-seq required preclinically.",
            "Cross-refs: Figure 8; Supplementary Figure S7; Methods §2.10.",
        ]),
    ]
    for title, bullets in sections:
        heading(doc, title, 2)
        for b in bullets:
            para(doc, b)


def _write_supplementary_figures(doc: Document, fig: dict[str, Path]) -> None:
    supp_legends = L.supplementary_figure_legends()
    for key in ["sfig1", "sfig2", "sfig3", "sfig4", "sfig5", "sfig6", "sfig7", "sfig8", "sfig9", "sfig10"]:
        if key in fig and key in supp_legends:
            fig_block(doc, fig[key], supp_legends[key], width=5.8)


def _write_supplementary_tables(doc: Document) -> None:
    supp_notes = L.supplementary_table_notes()

    table_block(doc, "Supplementary Table S1. Held-out guide validation.",
        ["Guide", "Editor", "Rank", "Percentile", "Pass", "padj", "Evidence"],
        [[g, D.held_out["validated_guide_recovery"]["per_guide"][g]["editor"],
          f"{D.held_out['validated_guide_recovery']['per_guide'][g]['rank']}/{D.held_out['validated_guide_recovery']['per_guide'][g]['n_editor_guides']}",
          f"{D.held_out['validated_guide_recovery']['per_guide'][g]['percentile']:.1f}%",
          "Yes" if D.held_out["validated_guide_recovery"]["per_guide"][g]["pass"] else "No",
          f"{D.held_out['validated_guide_recovery']['per_guide'][g]['padj']:.2e}",
          D.held_out["validated_guide_recovery"]["per_guide"][g]["evidence"]]
         for g in ["PWS_G.4670", "PWS_G.4363", "PWS_G.9414", "PWS_G.8738"]])
    table_legend_block(doc, supp_notes["S1"])

    table_block(doc, "Supplementary Table S2. Collateral imprinting expression (% WT, GSE243185).",
        ["Gene", "Tet1 %WT", "VP64 %WT"],
        [["SNRPN", "52.3", "246.3"], ["SNHG14", "99.6", "0.0"],
         ["UBE3A", "51.8", "47.5"], ["ATP10A", "65.2", "34.0"],
         ["NDN", "0.0", "0.06"], ["NPAP1", "57.1", "35.3"]])
    table_legend_block(doc, supp_notes["S2"])

    table_block(doc, "Supplementary Table S3. ENCODE tracks at PWS locus.",
        ["Track", "Items", "Status"],
        [[k, v["n_items"], v["status_code"]] for k, v in D.encode["tracks"].items()])
    table_legend_block(doc, (
        "Supplementary Table S3. ENCODE regulatory track retrieval summary (UCSC REST API, chr15:24.8–32.7 Mb).\n\n"
        "All 8 tracks returned HTTP 206 (partial content) with exactly 5,000 items each—the UCSC API per-request cap. "
        "This provides representative but not exhaustive regulatory coverage of the PWS locus.\n\n"
        "Inference: Brain-specific DNase and H3K27ac tracks confirm neuronal regulatory context at the locus; "
        "embryo tracks provide iPSC-differentiation proxy. Partial responses do not affect guide-level ranking "
        "but limit exhaustive CTCF/TFBS inventory.\n\n"
        "Cross-references: Results §1; Methods §2.5 (ENCODE); Figure 1."
    ))

    table_block(doc, "Supplementary Table S4. Dose optimization cross-method comparison.",
        ["Method", "w_Tet1", "w_VP64", "SNRPN%", "SNHG14%"],
        [["Differential evolution (G4670+G9414)", "0.987", "0.062", "70.0", "98.3"],
         ["Bayesian GP best EI (G4363+G9414)", f"{D.bayesian['gp_best_ei']['w_tet1']:.3f}",
          f"{D.bayesian['gp_best_ei']['w_vp64']:.3f}",
          f"{D.bayesian['gp_best_ei']['predicted_SNRPN_pct']:.1f}",
          f"{D.bayesian['gp_best_ei']['predicted_SNHG14_pct']:.1f}"]])
    table_legend_block(doc, (
        "Supplementary Table S4. Cross-method dose optimization comparison.\n\n"
        "Differential evolution (DE) on G4670+G9414: w_VP64=0.062, SNRPN=70.0% (window boundary). "
        f"Bayesian GP best EI on G4363+G9414: w_VP64={D.bayesian['gp_best_ei']['w_vp64']:.3f}, "
        f"SNRPN={D.bayesian['gp_best_ei']['predicted_SNRPN_pct']:.1f}% (mid-window). "
        "Dual-window GP optimum at w_VP64=0.123 yields 70.4%/70.1% (both at lower bound).\n\n"
        "Inference: Optimizer agreement on near-maximal Tet1 (w≈0.99) but discordance on VP64 fraction reflects "
        "guide-pair-specific methylation-validation vs screen-rank tradeoff. Primary catalog uses DE+G4670.\n\n"
        "Cross-references: Supplementary Figure S4; Figure 5; Methods §2.7."
    ))

    table_block(doc, "Supplementary Table S5. Top 15 Tet1 guides by rules_reactivation_score.",
        ["Rank", "Guide", "Score", "padj", "Dist CpG (bp)", "Meth % edited"],
        [[str(i+1), r.grna_id, f"{r.rules_reactivation_score:.4f}", f"{r.padj:.2e}",
          str(int(r.methylation_dist_bp)) if pd.notna(r.methylation_dist_bp) else "—",
          f"{r.methylation_pct_edited:.1f}" if pd.notna(r.methylation_pct_edited) else "—"]
         for i, r in enumerate(D.tet1_ranked.head(15).itertuples())])
    table_legend_block(doc, (
        "Supplementary Table S5. Top 15 dCas9-Tet1 guides ranked by rules_reactivation_score (n=57 significant Tet1 guides).\n\n"
        "PWS_G.4670 ranks first (score=1.02, padj=9.2×10⁻¹⁵, CpG distance=275 bp, edited methylation=37.2%). "
        "PWS_G.4363 is not in top 15 (rank 22; score=1.003, padj=2.95×10⁻⁶)—included in Supp. Table S1 as validation comparator. "
        "Top guides cluster within 350 bp of assayed ICR CpGs with methylation Δ≈−54.9%.\n\n"
        "Inference: The scoring function prioritizes screen significance and ICR proximity; experimental methylation "
        "validation provides independent confirmation not fully captured by padj alone.\n\n"
        "Cross-references: Figure 4; Results §2; Methods §2.6."
    ))

    rows6 = []
    for r in D.catalog.itertuples():
        rows6.append([str(int(r.catalog_rank)), r.strategy, r.tet1_grna_id,
                      r.vp64_grna_id if pd.notna(r.vp64_grna_id) else "—",
                      f"{r.predicted_SNRPN_pct_WT:.1f}", f"{r.predicted_SNHG14_pct_WT:.1f}",
                      f"{r.recommended_w_tet1:.3f}", f"{r.recommended_w_vp64:.3f}",
                      f"{r.uncertainty:.3f}", f"{r.genome_offtarget_risk:.2f}"])
    table_block(doc, "Supplementary Table S6. Complete therapeutic design catalog (n=25).",
        ["Rank", "Strategy", "Tet1", "VP64", "SNRPN%", "SNHG14%", "w_Tet1", "w_VP64", "Unc.", "Off-target"], rows6)
    table_legend_block(doc, supp_notes["S6"])

    table_block(doc, "Supplementary Table S7. Subregion guide counts and coordinates (GRCh38).",
        ["Subregion", "hg38 coordinates", "Significant guides"],
        [["pws_icr", "24,948,846–24,955,292", str(D.integration["subregion_hg38_distribution"]["pws_icr"])],
         ["snrpn_snh14", "—", str(D.integration["subregion_hg38_distribution"]["snrpn_snh14"])],
         ["snord116_cluster", "—", str(D.integration["subregion_hg38_distribution"]["snord116_cluster"])],
         ["pws_critical (other)", "—", str(D.integration["subregion_hg38_distribution"]["pws_critical"])]])
    table_legend_block(doc, (
        "Supplementary Table S7. PWS locus subregion definitions and significant guide counts (GRCh38).\n\n"
        "ICR spans ~6.4 kb and contains 80.4% of significant hits—defining the Tet1 therapeutic axis. "
        "GABAA cluster (28.5–31.0 Mb) contains zero top-catalog guide targets (>3.5 Mb from ICR).\n\n"
        "Cross-references: Figure 3; Methods §2.3; Supplementary Figure S1."
    ))

    table_block(doc, "Supplementary Table S8. Organoid fold-changes (GSE262700).",
        ["Gene", "Control", "Edited", "Fold-change"],
        [[r.gene_symbol, f"{r.control_organoid:.4f}", f"{r.edited_organoid:.4f}",
          f"{r.fold_change_edited_vs_control:.1f}"] for r in D.organoid_fc.itertuples()])
    table_legend_block(doc, (
        "Supplementary Table S8. Nemoto 2025 organoid gene-level fold-changes (edited vs control hypothalamic organoids).\n\n"
        "SNRPN: 5,297× from near-zero baseline (0.00038→2.007). SNHG14: 158× (0.034→5.372). "
        "UBE3A collateral: 1.05× (unchanged). MKRN3: 42.8× (unintended paternal gene reactivation—monitor in validation).\n\n"
        "Cross-references: Figure 7A; Supplementary Figure S6; Methods §2.4."
    ))

    table_block(doc, "Supplementary Table S9. Neuron vs iPSC concordance.",
        ["Editor:Gene", "iPSC %WT", "Neuron %WT", "Δ", "Concordant"],
        [[c["editor_gene"], c["ipsc_pct_WT"], c["neuron_pct_WT"], c["delta_pct"], c["direction_concordant"]]
         for c in D.neuron_conc["comparisons"]])
    table_legend_block(doc, (
        "Supplementary Table S9. Cell-type-specific expression concordance: hypothalamic neuron (GSE285305) vs iPSC bulk (GSE243185).\n\n"
        "Tet1 SNRPN: concordant direction (52.3%→47.2%, Δ=−5.2%). Tet1 SNHG14: concordant but attenuated in neurons "
        "(99.6%→56.8%, Δ=−42.9%). VP64 SNRPN: discordant (246.3%→25.7%)—VP64 dosing must be calibrated per cell type.\n\n"
        "Cross-references: Figure 7B; Supplementary Figure S9; Methods §2.9.3."
    ))

    table_block(doc, "Supplementary Table S10. Locus-level protospacer uniqueness.",
        ["Guide", "Protospacer", "Exact dup.", "1-mm seed", "Risk"],
        [["PWS_G.4670", D.locus_ot["guide_assessments"]["tet1_PWS_G.4670"]["protospacer"],
          D.locus_ot["guide_assessments"]["tet1_PWS_G.4670"]["exact_duplicates_in_locus"],
          D.locus_ot["guide_assessments"]["tet1_PWS_G.4670"]["near_seed_matches_1mm"],
          D.locus_ot["guide_assessments"]["tet1_PWS_G.4670"]["offtarget_locus_risk"]],
         ["PWS_G.9414", D.locus_ot["guide_assessments"]["vp64_PWS_G.9414"]["protospacer"],
          D.locus_ot["guide_assessments"]["vp64_PWS_G.9414"]["exact_duplicates_in_locus"],
          D.locus_ot["guide_assessments"]["vp64_PWS_G.9414"]["near_seed_matches_1mm"],
          D.locus_ot["guide_assessments"]["vp64_PWS_G.9414"]["offtarget_locus_risk"]]])

    table_legend_block(doc, (
        "Supplementary Table S10. Locus-level protospacer uniqueness within 464 significant Rohm 2025 screen sites.\n\n"
        "Both recommended guides: 0 exact 20-mer duplicates, 0 seed-region (1-mm) matches among screened sites. "
        "Scope limited to tiled library—not genome-wide.\n\n"
        "Inference: Low risk of competing on-target sites within the screen library; genome-wide assessment in Supp. Fig. S7.\n\n"
        "Cross-references: Figure 8; Methods §2.9.4, §2.10."
    ))

    table_block(doc, "Supplementary Table S11. Data and output file inventory.",
        ["File", "Description", "Main text reference"],
        [["pws_therapeutic_design_catalog.csv", "25 ranked designs", "Table 2, Supp. Table S6"],
         ["held_out_benchmark.json", "Guide recovery metrics", "Figure 6, Supp. Table S1"],
         ["organoid_concordance.json", "GSE262700 validation", "Figure 7, Supp. Table S8"],
         ["cas_offinder_offtarget.json", "Genome-wide off-target", "Figure 8, Supp. Fig. S7"],
         ["sensitivity_report.json", "Monte Carlo uncertainty", "Supp. Fig. S2"],
         ["experimental_validation_protocol.json", "Wet-lab protocol", "Results §5, Supp. Table S12"],
         ["gse28525_pws_probe_methylation.csv", "450K probe mapping", "Supp. Fig. S8"],
         ["forward_model_predictions.csv", "All guide scores", "Supp. Table S5"],
         ["pareto_therapeutic_designs.csv", "Pareto frontier", "Supp. Fig. S3"],
         ["circuit_ode_simulations.csv", "ODE scenarios", "Supp. Fig. S5"]])
    table_legend_block(doc, (
        "Supplementary Table S11. Complete inventory of primary output files with main-text cross-references.\n\n"
        "All files are generated by scripts/run_pipeline.py and scripts/run_phase12.py with UTC-stamped JSON audit trails. "
        "Machine-readable catalog (Supp. Table S6) should be deposited with supplementary data at NAR online.\n\n"
        "Cross-references: Data Availability; Methods §2.1."
    ))

    table_block(doc, "Supplementary Table S12. Experimental validation arms.",
        ["Arm", "Label", "Editors", "Priority"],
        [[a["arm_id"], a["label"], "; ".join(a["editors"]), a["priority"]]
         for a in D.protocol["test_arms"]])
    table_legend_block(doc, (
        "Supplementary Table S12. Prospective experimental validation protocol arms.\n\n"
        "Arm A (primary): hybrid PWS_G.4670+PWS_G.9414 at w_VP64≈0.06. Arm B (comparator): Tet1 monotherapy. "
        "Arm C: unedited PWS UPD15 control. Arm D: healthy WT iPSC reference. "
        f"Minimum n={D.protocol['sample_size_guidance']['minimum_per_arm']} biological replicates per arm. "
        f"Go criterion: both SNRPN and SNHG14 in 70–130% WT; no >2-fold collateral at UBE3A/GABRB3. "
        f"Iterate criterion: SNRPN 50–70% with SNHG14 in window—reduce VP64 and retest.\n\n"
        "Cross-references: Results §5; Methods §2.12; experimental_validation_protocol.json."
    ))


def _write_cross_reference_index(doc: Document) -> None:
    rows = [
        ["Results §1 (Digital twin)", "Fig. 1, 3", "Supp. Fig. S1", "Supp. Tables S3, S7, S11", "Methods S2–S4"],
        ["Results §2 (Forward model)", "Fig. 2, 4", "—", "Supp. Tables S2, S5", "Methods S5"],
        ["Results §3 (Hybrid optimization)", "Fig. 5", "Supp. Figs S2–S4, S10", "Tables 1, S4, S6", "Methods S6, S8"],
        ["Results §4 (Validation)", "Fig. 6, 7", "Supp. Figs S6, S8–S9", "Supp. Tables S1, S8–S9", "Methods S3, S7"],
        ["Results §5 (Protocol)", "—", "—", "Supp. Table S12", "protocol JSON"],
        ["Results §6 (Safety)", "Fig. 8", "Supp. Figs S7–S8", "Supp. Tables S2, S10", "Cas-OFFinder scan"],
        ["Discussion (Exploratory)", "—", "Supp. Fig. S5", "—", "Methods S9–S10"],
    ]
    table_block(doc, "Cross-reference index linking main Results sections to supplementary materials.",
        ["Results section", "Main figures", "Supplementary figures", "Supplementary tables", "Methods"],
        rows)


if __name__ == "__main__":
    import pandas as pd  # noqa: F401 — used in table builders
    paths = generate_all_figures()
    build_document(paths)

"""Extensive Materials and Methods section for NAR Methods manuscript."""
from __future__ import annotations

from docx import Document

import manuscript_data as D


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str) -> None:
    from docx.shared import Pt
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15


def table_block(doc: Document, caption: str, headers: list[str], rows: list[list]) -> None:
    from docx.shared import Pt
    cap = doc.add_paragraph()
    r = cap.add_run(caption)
    r.bold = True
    r.font.size = Pt(10)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def write_materials_and_methods(doc: Document) -> None:
    """Write full NAR-style Materials and Methods with systematic sub-headings."""

    heading(doc, "Materials and Methods")

    # ------------------------------------------------------------------
    heading(doc, "Study design and overview", 2)
    para(doc, (
        "We developed a computational design-prioritization platform for CRISPR epigenome-editing therapy in maternal "
        "UPD15 Prader–Willi syndrome (PWS). The study is a retrospective integrative analysis of publicly available multi-omic "
        "datasets with prospective in-silico therapeutic design, optimization, and validation. No new human subjects data were "
        "generated. The workflow comprises: (i) harmonized data acquisition on GRCh38; (ii) construction of an allele-aware "
        "PWS-locus digital twin; (iii) editor-stratified forward modeling; (iv) hybrid Tet1+VP64 dose optimization; "
        "(v) multi-layer retrospective validation; (vi) uncertainty quantification; and (vii) export of a ranked therapeutic "
        "catalog with experimental protocol. Pipeline orchestration: scripts/run_pipeline.py (phases 3, 5, 6, 9, 10, 11, 12) "
        "and scripts/run_phase12.py. All intermediate and final outputs are UTC-stamped JSON/CSV artifacts in data/models/."
    ))

    # ------------------------------------------------------------------
    heading(doc, "Ethics statement", 2)
    para(doc, (
        "This study used only publicly available, de-identified datasets deposited in NCBI GEO and the CRISPRepi database. "
        "Original human iPSC, organoid, and methylation data were generated under ethics approvals described in Rohm et al. "
        "(2025) Cell Genomics and Nemoto et al. (2025) Nature Communications. No additional ethics approval was required for "
        "computational reanalysis."
    ))

    # ------------------------------------------------------------------
    heading(doc, "Public datasets", 2)
    para(doc, (
        "Table 3 summarizes all primary datasets. Each accession was downloaded from NCBI GEO or the CRISPRepi web portal "
        "(https://crisprepi.com) and processed according to the subsections below. Genome coordinates were unified to "
        "GRCh38/hg38 unless otherwise noted."
    ))
    table_block(doc,
        "Table 3. Public datasets used in this study.",
        ["GEO/accession", "Reference", "Data type", "Samples / scale", "Role in platform"],
        [
            ["GSE285306", "Rohm 2025", "gRNA epigenome-editing screen", "23,506 editing records; 464 sig. sites", "Primary guide ranking and significance"],
            ["GSE243185", "Rohm 2025", "Bulk RNA-seq", "UPD15 iPSC edited vs control", "Editor-class expression anchors (Tet1/VP64)"],
            ["GSE285300", "Rohm 2025", "Bisulfite sequencing", "ICR methylation amplicons", "Methylation validation; G4363 comparator"],
            ["GSE285305", "Rohm 2025", "Bulk RNA-seq", "Hypothalamic neuron derivatives", "Neuron-specific concordance (Supp. Fig. S9)"],
            ["GSE262700", "Nemoto 2025", "Organoid scRNA-seq", "PWS hypothalamic organoids", "Organoid reactivation validation"],
            ["GSE152098", "Cousminer 2021", "ATAC-seq bigWig (RAW)", "14 bigWig files; 3 cell types", "ICR accessibility quantification"],
            ["GSE28525", "Imprinting reference", "450K methylation array", "71 probes in PWS window", "Imprinted DMR collateral map"],
            ["GSE298378", "Methylation array", "450K beta values", "1,544 probes; 8 samples", "Extended probe coverage"],
            ["CRISPRepi", "Shi 2025 NAR", "Epigenome-editing database", "2,520 human records", "Editor-class mechanism priors"],
            ["ENCODE (UCSC API)", "ENCODE 2020", "Regulatory annotations", "8 tracks × 5,000 features", "DNase, H3K27ac, H3K4me3, CTCF, TFBS"],
            ["UCSC hg38", "Genome Reference", "Reference assembly", "25 chromosomes", "Coordinate system; Cas-OFFinder scan"],
        ],
    )
    para(doc, (
        "Therapeutic expression targets were defined relative to wild-type iPSC mean counts from GSE243185: SNRPN WT mean "
        "231.8 counts; SNHG14 WT mean 23,127.5 counts. The therapeutic adequacy window was predefined as 70–130% of WT "
        "expression for both genes, representing paternal-equivalent dual-gene rescue."
    ))

    # ------------------------------------------------------------------
    heading(doc, "Genome reference, coordinate harmonization, and locus definition", 2)
    heading(doc, "Reference genome", 3)
    para(doc, (
        "All analyses used GRCh38 (hg38) as the primary reference assembly. Chromosome sequences were obtained from "
        "UCSC hgdownload (https://hgdownload.soe.ucsc.edu/goldenPath/hg38/chromosomes/) and cached locally for "
        "25 chromosomes (chr1–22, chrX, chrY, chrM) in data/genome/hg38/."
    ))
    heading(doc, "liftOver and base-pair integration", 3)
    para(doc, (
        "Rohm et al. (2025) screen coordinates were originally reported on hg19. All 23,506 editing records were lifted "
        "to hg38 using UCSC liftOver chain files (mapping rate: 100%; 23,506/23,506). Example: top Tet1 hit PWS_G.4670 "
        "maps from hg19 chr15:25,200,068 to hg38 chr15:24,954,921 (~246 kb coordinate shift on chr15). Significant sites "
        "(padj < 0.05) were merged at base-pair resolution yielding 464 unique gRNA-target sites. Integration report: "
        "data/integrated/integration_report.json."
    ))
    heading(doc, "PWS locus boundaries and subregions", 3)
    para(doc, (
        f"The PWS/Angelman critical region was defined as chr15:24,800,000–32,700,000 (GRCh38), containing "
        f"{D.locus_summary['gene_count']} UCSC-annotated genes and {D.locus_summary['cpg_island_count']} CpG islands. "
        "Functional subregions (Supplementary Table S7): PWS-ICR (hg38 chr15:24,948,846–24,955,292; ~6.4 kb); "
        "SNRPN–SNHG14 transcriptional unit (chr15:24,750,000–26,300,000); SNORD116 cluster (chr15:26,300,000–28,500,000); "
        "GABAA receptor gene cluster (chr15:28,500,000–31,000,000; preservation priority for collateral safety). "
        "Key paternally expressed genes: SNRPN, SNHG14, SNORD116, MKRN3, MAGEL2, NDN, NPAP1. "
        "Key maternally expressed genes: UBE3A, ATP10A."
    ))

    # ------------------------------------------------------------------
    heading(doc, "Data preprocessing", 2)
    heading(doc, "CRISPR epigenome-editing screen (GSE285306)", 3)
    para(doc, (
        "Screen records were curated from Rohm et al. (2025) supplementary data and CRISPRepi cross-reference. Each record "
        "contains: gRNA identifier (PWS_G.xxxx nomenclature), protospacer sequence (20 nt), editor system "
        "(dCas9-Tet1, dCas9-VP64, dCas9-KRAB, dCas9-p300), genomic coordinates, screen significance (padj), and overlapping "
        "annotation flags. Multiple testing correction (padj) was applied by Rohm et al.; we retained padj < 0.05 as the "
        "significance threshold (464 sites). Editor-class counts among significant sites: 57 dCas9-Tet1, 17 dCas9-VP64, "
        "with additional KRAB/p300 records excluded from therapeutic forward modeling due to silencing/activation "
        "mechanisms incompatible with PWS reactivation goals."
    ))
    heading(doc, "Bulk and neuron RNA-seq (GSE243185, GSE285305)", 3)
    para(doc, (
        "Editor-class mean expression counts were extracted for SNRPN, SNHG14, and collateral genes (UBE3A, ATP10A, NDN, "
        "MKRN3, NPAP1). Percent-of-WT was computed as (edited_mean / WT_mean) × 100. Anchors enforced in forward model: "
        "Tet1 SNRPN=52.336%, SNHG14=99.629%; VP64 SNRPN=246.298%, SNHG14=0%. Neuron-specific values (GSE285305): "
        "Tet1 SNRPN=47.2%, SNHG14=56.8%; VP64 SNRPN=25.7%—used for concordance validation only, not model training."
    ))
    heading(doc, "Organoid scRNA-seq (GSE262700)", 3)
    para(doc, (
        "Nemoto et al. (2025) PWS UPD15 hypothalamic organoid data were processed to compute fold-changes (edited vs control) "
        "for PWS-region genes. Key values: SNRPN 5,297× (control mean 0.00038→edited 2.007 counts); SNHG14 158× "
        "(0.034→5.372); percent cells expressing SNRPN in edited organoids 58.5%; SNHG14 73.6%. Editor construct: "
        "SunTag-dCas9-TET1 (same TET1 demethylation mechanism class as dCas9-Tet1 but different delivery format)."
    ))
    heading(doc, "Hypothalamic ATAC-seq bigWig reprocessing (GSE152098)", 3)
    para(doc, (
        "Fourteen bigWig files were extracted from GSE152098_RAW.tar and quantified using UCSC bigWigSummary (Linux binary "
        "via WSL) at five predefined regions: pws_critical, pws_icr, snrpn_snh14, snord116_cluster, gabaa_cluster. "
        f"ICR mean signal: Human ESCs 0.0573; hypothalamic progenitors 0.3317; hypothalamic neurons 0.6026 "
        f"(neuron/ESC fold={D.gse152098['neuron_vs_esc_icr_fold']:.2f}). Output: data/curated/gse152098_bigwig_regions.parquet. "
        "Script: scripts/reprocess_gse152098_raw.py."
    ))
    heading(doc, "450K methylation array mapping (GSE28525, GSE298378)", 3)
    para(doc, (
        "The Illumina HumanMethylation450 manifest (15,259 chr15 probes) was downloaded from UCSC ENCODE supplemental files. "
        "Probe coordinates were mapped to hg38 via liftOver. GSE28525: 71 probes in PWS window, 10 imprinted DMRs, 3 ICR-associated "
        "maternally methylated probes. GSE298378: 1,544 probes, 8 samples, beta range 0.009–0.980. Imprinted DMR classification "
        "used delta_mUPD and maternally_methylated flags from reference sample comparisons. "
        "Script: scripts/map_methylation_dmrs.py."
    ))
    heading(doc, "CRISPRepi mechanism priors", 3)
    para(doc, (
        f"CRISPRepi human records (n={D.crisprepi['n_human_records']}) were aggregated by editor class. dCas9-Tet1: 16 records, "
        f"62.5% report demethylation as primary epigenetic effect. dCas9-VP64: 134 records, 60.4% report transcriptional activation. "
        f"dCas9-KRAB: 1,489 records (not used for PWS therapeutic ranking). These priors inform editor-class confidence weights "
        "in the forward model but do not provide locus-specific PWS outcome predictions. Chr15 CRISPRepi coordinate hits: 43."
    ))
    heading(doc, "ENCODE regulatory tracks (UCSC REST API)", 3)
    para(doc, (
        "Regulatory features were retrieved from https://api.genome.ucsc.edu for chr15:24,800,000–32,700,000. "
        "Eight tracks succeeded (8/8): encRegTfbsClustered, wgEncodeReg4DnaseAllBrain, wgEncodeReg4DnaseAllEmbryo, "
        "wgEncodeReg4MarkH3k27acAllBrain, wgEncodeReg4MarkH3k27acAllEmbryo, wgEncodeReg4MarkH3k4me3AllBrain, "
        "wgEncodeReg4MarkH3k4me3AllEmbryo, wgEncodeReg4MarkCtcfAllBrain. API limitation: maximum 5,000 items per track "
        "(HTTP 206 partial content). Script: scripts/download_encode_reference.py."
    ))

    # ------------------------------------------------------------------
    heading(doc, "Allele-aware PWS-locus digital twin", 2)
    para(doc, (
        f"The digital twin is a heterogeneous graph integrating {sum(D.integration['node_counts'].values())} nodes and "
        f"{D.integration['edge_count']} edges (data/integrated/locus_merged_hg38.csv). Node types: "
        f"897 ATAC peaks (GSE152090 hypothalamic neuron atlas), 464 significant gRNA targets, 153 genes, 81 CpG islands, "
        f"80 bisulfite-assayed methylation CpGs. Edge types: (i) gRNA→CpG island if distance ≤500 bp; "
        "(ii) gRNA→gene if overlap; (iii) gene→gene adjacency within subregion; (iv) gRNA→ATAC peak proximity. "
        f"Subregion assignment: {D.integration['subregion_hg38_distribution']}. "
        "The twin is allele-aware in the sense that all coordinates and methylation states reflect the maternal UPD15 "
        "epigenotype context (biallelic maternal methylation at ICR) used in Rohm/Nemoto experimental systems."
    ))

    # ------------------------------------------------------------------
    heading(doc, "Editor-stratified forward model", 2)
    heading(doc, "Rationale for editor stratification", 3)
    para(doc, (
        "Tet1 (TET1 catalytic domain fusion) mediates active DNA demethylation at CpG dinucleotides; VP64 (VP64 transcriptional "
        "activation domain) enhances promoter transcription without demethylating the ICR. These mechanisms are not exchangeable. "
        "We attempted cross-editor machine learning (pooling all significant guides) and rejected it when 5-fold cross-validated "
        f"R²<0 (empirically confirmed: Tet1 within-editor Ridge CV R²={D.forward['tet1_within_editor_model']['cv_r2_mean']:.2f}). "
        "All subsequent ranking uses editor-stratified models."
    ))
    heading(doc, "rules_reactivation_score (primary)", 3)
    para(doc, (
        "The primary ranking function rules_reactivation_score integrates: (i) screen significance as −log10(padj); "
        "(ii) Tet1-specific distance to ICR anchor (log-transformed; Spearman ρ=−0.86 with score, n=57); "
        "(iii) methylation change at nearest assayed CpG (Spearman ρ=0.664, n=48); (iv) subregion classification bonus "
        "(pws_icr > snrpn_snh14); (v) editor-class bulk RNA priors from GSE243185. VP64 guides are ranked by screen significance "
        "and SNRPN promoter proximity within the snrpn_snh14 subregion. Output: data/models/forward_model_predictions.csv "
        "(464 site-level predictions) and data/models/therapeutic_candidates_ranked.csv (top Tet1 candidates)."
    ))
    heading(doc, "Supplementary within-editor Ridge regression", 3)
    para(doc, (
        f"A Ridge regression was trained on Tet1 guides only (n=57) with features log(distance to ICR) and SNHG14 targeting "
        f"flag. 5-fold CV R²={D.forward['tet1_within_editor_model']['cv_r2_mean']:.2f} "
        f"(±{D.forward['tet1_within_editor_model']['cv_r2_std']:.2f}). This model is reported for transparency but not used "
        "for catalog ranking due to negative predictive performance."
    ))

    # ------------------------------------------------------------------
    heading(doc, "Hybrid therapeutic design optimization", 2)
    heading(doc, "Phenomenological dose-response model", 3)
    para(doc, (
        "Dual-editor outcomes were modeled phenomenologically from Rohm 2025 bulk RNA editor-class anchors: "
        "SNRPN_pred(w_VP64) = (1 − w_VP64) × 52.3 + w_VP64 × 246.3; "
        "SNHG14_pred(w_VP64) = (1 − w_VP64) × 99.6 + w_VP64 × 0.0, "
        "where w_VP64 ∈ [0,1] is the relative VP64 dose fraction and w_Tet1 = 1 − w_VP64 (Tet1-dominated normalization). "
        "This assumes linear interpolation between editor-class outcomes—a deliberate simplification requiring experimental calibration."
    ))
    heading(doc, "Differential evolution optimizer", 3)
    para(doc, (
        f"Differential evolution (scipy.optimize.differential_evolution) searched (w_Tet1, w_VP64) for each Tet1×VP64 guide pair "
        f"from ranked candidate lists. Objective: maximize binary window scores (1 if 70≤%WT≤130, else 0) for SNRPN and SNHG14, "
        f"minus collateral penalty for GABAA cluster proximity. Total designs evaluated: "
        f"{D.optimization['summary']['n_designs_evaluated']}; dual-window designs: "
        f"{D.optimization['summary']['n_both_genes_in_target_window']} ({100*D.optimization['summary']['n_both_genes_in_target_window']/D.optimization['summary']['n_designs_evaluated']:.1f}%). "
        "Consistent optimum: w_Tet1≈0.987, w_VP64≈0.062. Output: data/models/optimization/optimized_therapeutic_designs.csv."
    ))
    heading(doc, "Pareto and Bayesian GP supplementary optimizers", 3)
    para(doc, (
        f"Non-dominated sorting identified {D.pareto['n_pareto_optimal']} Pareto-optimal designs across four objectives "
        "(SNRPN window, SNHG14 window, safety proxy, confidence). Bayesian Gaussian Process optimization (Matern kernel, "
        f"Expected Improvement; 25 initial samples) was run on G4363+G9414: best EI at w_Tet1={D.bayesian['gp_best_ei']['w_tet1']:.3f}, "
        f"w_VP64={D.bayesian['gp_best_ei']['w_vp64']:.3f}. These supplementary analyses are reported but differential evolution "
        "on G4670+G9414 is the primary optimization result."
    ))

    # ------------------------------------------------------------------
    heading(doc, "Retrospective validation suite", 2)
    heading(doc, "Held-out guide benchmark", 3)
    para(doc, (
        "Four experimentally validated guides (PWS_G.4670, G4363, G9414, G8738) were held out from model training and ranked "
        "post-hoc within editor class. Pass criterion: within-editor percentile ≥90 (top decile). "
        f"Results: 3/4 pass (75%). Tet1 Spearman ρ=0.401 (p=0.002, n=57); VP64 ρ=0.630 (p=0.007, n=17). "
        "Script: scripts/held_out_benchmark.py (or equivalent validation module)."
    ))
    heading(doc, "Organoid concordance", 3)
    para(doc, (
        "Direction concordance between GSE262700 organoid fold-changes and GSE243185 bulk %WT was assessed for SNRPN and SNHG14. "
        "Both genes: direction concordant (up in both datasets). Magnitude comparison explicitly not performed due to scale mismatch "
        "(organoid FC from near-zero baseline vs bulk %WT from healthy iPSC denominator)."
    ))
    heading(doc, "Neuron-iPSC concordance", 3)
    para(doc, (
        "GSE285305 neuron RNA compared to GSE243185 iPSC bulk for editor-class outcomes. Tet1: direction concordant for SNRPN "
        "(Δ=−5.2%) and SNHG14 (Δ=−42.9%). VP64 SNRPN: discordant magnitude (Δ=−220.6%). Overall neuron concordance test: fail "
        "on VP64 magnitude; pass on Tet1 direction."
    ))
    heading(doc, "Locus-level off-target uniqueness", 3)
    para(doc, (
        "Among 464 significant screen sites, exact protospacer duplicates and 1-mm seed matches were counted for recommended guides. "
        "PWS_G.4670 and PWS_G.9414: zero exact duplicates, zero 1-mm seed matches, low locus risk."
    ))

    # ------------------------------------------------------------------
    heading(doc, "Genome-wide off-target assessment", 2)
    para(doc, (
        "Cas-OFFinder-compatible in-silico scanning (scripts/cas_offinder_scan.py, scripts/offtarget_core.py) searched full hg38 "
        "for SpCas9 NGG PAM sites with 0–4 mismatches (mm0–mm4). Mismatch weights follow Cas-OFFinder convention. "
        f"PWS_G.4670: score={D.cas_off['per_guide']['PWS_G.4670']['cas_offinder_score']:.3f}, risk=low. "
        f"PWS_G.9414: score={D.cas_off['per_guide']['PWS_G.9414']['cas_offinder_score']:.3f}, risk=low. "
        "Limitations: no RNA/DNA bulge modeling; no epigenetic modulators; in-silico only. Biochemical validation "
        "(CHANGE-seq/GUIDE-seq) required before clinical translation."
    ))

    # ------------------------------------------------------------------
    heading(doc, "Uncertainty quantification", 2)
    para(doc, (
        f"Monte Carlo sensitivity analysis (scripts/sensitivity_analysis.py) drew {D.sensitivity['n_samples']:,} Gaussian "
        f"perturbations per catalog design with σ_pct = 25 × design uncertainty (rank 1: uncertainty=0.147, σ=3.68%). "
        f"Window probabilities computed for SNRPN, SNHG14, and joint (both in 70–130% window). Top design: "
        f"p_SNRPN={D.sensitivity['top_design']['p_SNRPN_in_window']:.3f}, p_SNHG14={D.sensitivity['top_design']['p_SNHG14_in_window']:.1f}, "
        f"p_both={D.sensitivity['top_design']['p_both_in_window']:.3f}. Uncertainty is heuristic (guide confidence), not a "
        "calibrated Bayesian posterior from biological replicates."
    ))

    # ------------------------------------------------------------------
    heading(doc, "Exploratory analyses (not used for catalog ranking)", 2)
    para(doc, (
        f"Graph propagation: Personalized PageRank (α={D.graph['damping_alpha']}) on {D.graph['n_nodes']} nodes—SNRPN propagated "
        f"scores ≈0 for validated guides; limited discriminative utility. ODE circuit model: three-node AgRP/POMC/OXT linear "
        f"relaxation; hybrid hyperphagia_proxy=0.842 vs Tet1 mono 0.760. CRISPRepi transfer priors: mechanism-class confidence only. "
        "These are reported in Supplementary Methods S9–S10 and Supplementary Figure S5."
    ))

    # ------------------------------------------------------------------
    heading(doc, "Therapeutic catalog assembly and experimental protocol", 2)
    para(doc, (
        "Final catalog (scripts/build_final_catalog.py): 25 ranked designs merged from optimization, validation scores, "
        "off-target risk, and uncertainty. Top design: Hybrid PWS_G.4670+PWS_G.9414. Experimental protocol "
        "(data/models/experimental_protocol/experimental_validation_protocol.json): four arms (A: hybrid, B: Tet1 mono, "
        "C: unedited PWS, D: healthy WT); primary endpoints ICR bisulfite (~40% methylation target), SNRPN/SNHG14 qPCR/RNA-seq "
        "(70–130% WT), collateral panel; minimum n=3 biological replicates per arm."
    ))

    # ------------------------------------------------------------------
    heading(doc, "Software and computational environment", 2)
    para(doc, (
        "Python 3.13. Core packages: pandas 2.x, numpy, scipy (differential_evolution, Spearman), scikit-learn (Ridge, CV), "
        "matplotlib (figures). External tools: UCSC liftOver, bigWigSummary (WSL Linux binary), UCSC REST API. "
        "Pipeline scripts in scripts/; reproducible execution via scripts/run_pipeline.py --phase [3|5|6|9|10|11|12|all]. "
        "Code and processed data: [repository URL to be inserted at submission]."
    ))

    # ------------------------------------------------------------------
    heading(doc, "Statistical analysis", 2)
    para(doc, (
        "Spearman rank correlation (scipy.stats.spearmanr) for guide ranking concordance with two-sided p-values. "
        "Screen significance: padj < 0.05 (Benjamini-Hochberg correction applied by Rohm et al.). Monte Carlo: 50,000 samples "
        "per design; Gaussian perturbation model. Pareto optimality: non-dominated sorting across four objectives. "
        "No formal multiple-testing correction applied across our retrospective validation layers (exploratory concordance framework). "
        "Experimental protocol specifies n≥3 per arm with post-hoc power calculation after pilot variance estimation."
    ))
